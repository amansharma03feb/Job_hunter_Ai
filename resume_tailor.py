"""
Resume & Outreach Builder — no Claude API calls, zero tailoring cost.

Generates:
  - Base resume DOCX (from RESUME_FULL in config)
  - Template cover letter DOCX
  - Template email subject + body

This replaces Claude-based tailoring to save API credits.
Claude is only used for ATS scoring (already capped at 15 calls/run).
"""

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import RESUME_FULL, SENDER_NAME

OUTPUT_DIR  = "output"
RESUMES_DIR = os.path.join(OUTPUT_DIR, "resumes")
CL_DIR      = os.path.join(OUTPUT_DIR, "cover_letters")


def _safe_filename(s):
    return re.sub(r"[^a-zA-Z0-9_\-]", "_", str(s or ""))[:40]


# ── Base resume DOCX (no Claude) ──────────────────────────────────────────────

def save_resume_docx(resume_text, job_title, company):
    """Save RESUME_FULL as a formatted DOCX. resume_text param ignored (use base)."""
    os.makedirs(RESUMES_DIR, exist_ok=True)
    date_str = datetime.now().strftime("%Y%m%d")
    fname = f"Aman_Sharma_Resume_{_safe_filename(job_title)}_{_safe_filename(company)}_{date_str}.docx"
    path = os.path.join(RESUMES_DIR, fname)

    doc = Document()
    # Name header
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(SENDER_NAME)
    run.bold = True
    run.font.size = Pt(16)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("amansharma03feb@gmail.com  |  linkedin.com/in/amansharma03feb  |  India (open to relocate: Ireland/UK/EU)")

    doc.add_paragraph()

    for line in RESUME_FULL.strip().split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        if stripped.isupper() or (stripped.endswith(":") and len(stripped) < 40):
            r = p.add_run(stripped)
            r.bold = True
        else:
            p.add_run(stripped)

    doc.save(path)
    return path


# ── Template cover letter DOCX (no Claude) ────────────────────────────────────

def save_cover_letter_docx(cl_text, job_title, company, hr_name):
    """Generate a template cover letter DOCX. cl_text param ignored (use template)."""
    os.makedirs(CL_DIR, exist_ok=True)
    date_str = datetime.now().strftime("%Y%m%d")
    fname = f"CoverLetter_{_safe_filename(job_title)}_{_safe_filename(company)}_{date_str}.docx"
    path = os.path.join(CL_DIR, fname)

    greeting_name = str(hr_name or "").split()[0] if hr_name else "Hiring Manager"

    body = f"""Dear {greeting_name},

I am writing to express my strong interest in the {job_title} position at {company}.

With 6+ years of experience delivering healthcare data platforms, AI-enabled products, and enterprise ETL/MDM solutions, I bring a proven track record across US healthcare, fintech, and SaaS domains. At CloudAngles, I designed and owned a Member Identity Resolution & MDM platform for a Fortune-class US health insurer — architecting dual-mode ingestion via Kafka (real-time) and Airflow (batch) on AWS, cutting data processing time by 50%, and governing Snowflake data structures for Power BI, MicroStrategy, and Tableau reporting layers.

I am actively seeking Senior BA / Technical Product Owner roles in Ireland, UK, and global teams with travel exposure, with a relocation timeline of July 2026. I believe my background closely aligns with what {company} is looking for, and I would welcome the opportunity to discuss how I can contribute to your team.

Please find my resume attached. I look forward to hearing from you.

Warm regards,
Aman Sharma
Senior Business Analyst
amansharma03feb@gmail.com
LinkedIn: linkedin.com/in/amansharma03feb"""

    doc = Document()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h.add_run(f"{SENDER_NAME}\namansharma03feb@gmail.com\n{datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph()

    for para in body.strip().split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.save(path)
    return path


# ── Template email (no Claude) ────────────────────────────────────────────────

def draft_email(job_title, company, hr_first_name):
    """Return (subject, body) using a fixed template. Zero Claude cost."""
    first = str(hr_first_name or "").strip() or "Hiring Manager"

    subject = f"Application: {job_title} — Aman Sharma, Senior BA"

    body = f"""Dear {first},

I came across the {job_title} opening at {company} and would love to be considered.

I'm a Senior Business Analyst with 6+ years in healthcare data (MDM, Kafka/Airflow ETL, Snowflake, HIPAA) and AI/LLM delivery. I recently led the Member Identity Resolution platform for a Fortune-class US health insurer and have delivered multiple 5-star client engagements. I hold a CSPO certification and am actively targeting Ireland, UK, and global roles with a July 2026 relocation timeline.

My tailored resume and cover letter are attached. Happy to connect for a quick call.

Best regards,
Aman Sharma
amansharma03feb@gmail.com | linkedin.com/in/amansharma03feb"""

    return subject, body


# ── Kept for backward compatibility — not used ─────────────────────────────────

def tailor_resume_text(job_title, company, jd_text):
    """No-op — base resume used instead of Claude tailoring."""
    return None


def generate_cover_letter(job_title, company, hr_name, jd_text):
    """No-op — template cover letter used instead of Claude generation."""
    return None
